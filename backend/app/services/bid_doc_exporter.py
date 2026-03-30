"""
投标文件 Word 导出服务 — 将章节内容渲染为格式化的 .docx 文件

输出结构:
  封面 → 目录 → 9 章正文 → 报价表 → 资质清单
  （template/credential 类型章节输出占位提示，ai/manual 章节渲染正文）

架构红线:
  - 报价数值来自 QuotationSheet，禁止用 LLM 输出
  - tenant_id 隔离
"""
import os
import re
from datetime import datetime
from typing import Optional

from docx import Document
from docx.shared import Pt, Cm, RGBColor, Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.enum.table import WD_TABLE_ALIGNMENT

from sqlalchemy import select
from sqlalchemy.ext.asyncio import AsyncSession

from app.models.bid_project import BidProject, BidChapter
from app.models.enterprise import Enterprise
from app.models.credential import Credential
from app.models.quotation import QuotationSheet, QuotationItem
from app.services.bid_project_service import BidProjectService

# ---------- 中文数字映射 ----------
_CN_NUMS = {
    1: "一", 2: "二", 3: "三", 4: "四", 5: "五",
    6: "六", 7: "七", 8: "八", 9: "九", 10: "十",
}

# ---------- 报价品类中文 ----------
_FOOD_CATEGORY_LABELS = {
    "vegetable": "蔬菜类", "meat": "肉类", "seafood": "水产类",
    "egg_poultry": "蛋禽类", "dry_goods": "干货类", "condiment": "调料类",
}

# ---------- 资质类型中文 ----------
_CRED_TYPE_LABELS = {
    "food_license": "食品经营许可证", "business_license": "营业执照",
    "haccp": "HACCP认证", "iso22000": "ISO22000认证", "sc": "SC认证",
    "animal_quarantine": "动物防疫合格证", "cold_chain_transport": "冷链运输资质",
    "health_certificate": "从业人员健康证", "liability_insurance": "公众责任险",
    "quality_inspection": "质量检验报告", "organic_cert": "有机认证",
    "green_food": "绿色食品认证", "performance": "业绩证明",
    "award": "荣誉证书", "other": "其他",
}

# ---------- 正文智能排版正则 ----------
_SECTION_TITLE_RE = re.compile(r'^第[一二三四五六七八九十\d]+[节章部分][\s\u4e00-\u9fff]')
_NUMBERED_ITEM_RE = re.compile(r'^[一二三四五六七八九十]+[、．.]')
_CLAUSE_TITLE_RE = re.compile(r'^第[一二三四五六七八九十百\d]+条')
_SUB_ITEM_RE = re.compile(r'^(\d{1,2}[\.、]|（[一二三四五六七八九十]）)')
_SUB_CLAUSE_RE = re.compile(r'^(（\d+）|\(\d+\)|[①②③④⑤⑥⑦⑧⑨⑩])')
_HEADING_RE = re.compile(r'^#{1,4}\s+')


class BidDocExporter:
    """投标文件 Word 导出器"""

    def __init__(self, session: AsyncSession):
        self.session = session

    async def export(
        self, project_id: int, tenant_id: int
    ) -> str:
        """
        导出投标文件为 Word。

        Returns:
            生成的 .docx 文件绝对路径
        """
        # 1. 加载数据
        svc = BidProjectService(self.session)
        project = await svc.get_project(project_id, tenant_id)
        if not project:
            raise ValueError("投标项目不存在")

        if not project.chapters:
            raise ValueError("项目尚未初始化章节，请先生成章节内容")

        enterprise = await self._load_enterprise(project.enterprise_id, tenant_id)
        credentials = await self._load_credentials(project.enterprise_id, tenant_id) if enterprise else []
        quotation = await self._load_latest_quotation(project_id, tenant_id)

        # 2. 渲染 Word
        file_path = self._render_docx(project, enterprise, credentials, quotation)

        # 3. 更新项目的投标文件路径
        project.bid_doc_path = file_path
        await self.session.commit()

        return file_path

    # ========== 数据加载 ==========

    async def _load_enterprise(self, enterprise_id: Optional[int], tenant_id: int) -> Optional[Enterprise]:
        """加载企业信息 — 强制 tenant_id 隔离，防止跨租户越权"""
        if not enterprise_id:
            return None
        result = await self.session.execute(
            select(Enterprise).where(
                Enterprise.id == enterprise_id,
                Enterprise.tenant_id == tenant_id,  # 安全: 必须属于当前租户
            )
        )
        return result.scalar_one_or_none()

    async def _load_credentials(self, enterprise_id: Optional[int], tenant_id: int) -> list[Credential]:
        if not enterprise_id:
            return []
        result = await self.session.execute(
            select(Credential).where(
                Credential.enterprise_id == enterprise_id,
                Credential.tenant_id == tenant_id,
            ).order_by(Credential.cred_type)
        )
        return list(result.scalars().all())

    async def _load_latest_quotation(self, project_id: int, tenant_id: int) -> Optional[QuotationSheet]:
        result = await self.session.execute(
            select(QuotationSheet)
            .where(QuotationSheet.project_id == project_id)
            .where(QuotationSheet.tenant_id == tenant_id)
            .order_by(QuotationSheet.version.desc())
            .limit(1)
        )
        return result.scalar_one_or_none()

    # ========== Word 渲染 ==========

    def _render_docx(
        self,
        project: BidProject,
        enterprise: Optional[Enterprise],
        credentials: list[Credential],
        quotation: Optional[QuotationSheet],
    ) -> str:
        doc = Document()
        self._setup_styles(doc)

        # 封面
        self._render_cover(doc, project, enterprise)
        doc.add_page_break()

        # 目录页
        self._render_toc(doc, project.chapters)
        doc.add_page_break()

        # 正文章节
        sorted_chapters = sorted(project.chapters, key=lambda c: c.sort_order)
        for ch in sorted_chapters:
            self._render_chapter(doc, ch, enterprise, credentials, quotation)

        # 保存
        output_dir = os.path.join(
            os.path.dirname(os.path.dirname(os.path.dirname(os.path.abspath(__file__)))),
            "storage", "bid_outputs",
        )
        os.makedirs(output_dir, exist_ok=True)

        safe_name = project.project_name.replace("/", "_").replace("\\", "_").replace(" ", "_")[:60]
        filename = f"{safe_name}_投标文件_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
        file_path = os.path.join(output_dir, filename)

        doc.save(file_path)
        return file_path

    def _setup_styles(self, doc: Document):
        """设置全局文档样式"""
        style = doc.styles["Normal"]
        style.font.name = "宋体"
        style.font.size = Pt(12)
        style.paragraph_format.line_spacing = Pt(22)

        # 调整页边距
        for section in doc.sections:
            section.top_margin = Cm(2.5)
            section.bottom_margin = Cm(2.5)
            section.left_margin = Cm(3.0)
            section.right_margin = Cm(2.5)

    def _render_cover(self, doc: Document, project: BidProject, enterprise: Optional[Enterprise]):
        """渲染封面"""
        for _ in range(5):
            doc.add_paragraph()

        # 项目名称
        p = doc.add_paragraph()
        p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        run = p.add_run(project.project_name)
        run.font.size = Pt(22)
        run.font.bold = True

        doc.add_paragraph()

        # 副标题
        p = doc.add_paragraph()
        p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        run = p.add_run("投 标 文 件")
        run.font.size = Pt(26)
        run.font.bold = True

        for _ in range(3):
            doc.add_paragraph()

        # 元信息
        meta_items = []
        if enterprise:
            meta_items.append(f"投标人：{enterprise.name}")
        if project.tender_org:
            meta_items.append(f"招标人：{project.tender_org}")
        meta_items.append(f"编制日期：{datetime.now().strftime('%Y年%m月%d日')}")

        for item in meta_items:
            p = doc.add_paragraph()
            p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            run = p.add_run(item)
            run.font.size = Pt(14)

    def _render_toc(self, doc: Document, chapters: list[BidChapter]):
        """渲染目录页"""
        p = doc.add_paragraph()
        p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        run = p.add_run("目    录")
        run.font.size = Pt(18)
        run.font.bold = True

        doc.add_paragraph()

        sorted_chapters = sorted(chapters, key=lambda c: c.sort_order)
        for ch in sorted_chapters:
            p = doc.add_paragraph()
            run = p.add_run(f"{ch.chapter_no}  {ch.title}")
            run.font.size = Pt(14)
            p.paragraph_format.space_after = Pt(4)

    def _render_chapter(
        self,
        doc: Document,
        chapter: BidChapter,
        enterprise: Optional[Enterprise],
        credentials: list[Credential],
        quotation: Optional[QuotationSheet],
    ):
        """渲染单个章节"""
        # 章节标题
        doc.add_heading(f"{chapter.chapter_no}  {chapter.title}", level=1)

        # 根据 source 类型处理
        if chapter.source == "credential":
            self._render_credential_chapter(doc, chapter, enterprise, credentials)
        elif chapter.chapter_no == "第八章" and quotation:
            self._render_quotation_chapter(doc, chapter, quotation)
        elif chapter.content:
            self._render_content(doc, chapter.content)
        else:
            p = doc.add_paragraph()
            run = p.add_run("（本章节内容待补充）")
            run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)
            run.font.italic = True

    def _render_credential_chapter(
        self,
        doc: Document,
        chapter: BidChapter,
        enterprise: Optional[Enterprise],
        credentials: list[Credential],
    ):
        """渲染资质/企业简介类章节"""
        # 如果有 AI/手动编辑的内容，优先输出
        if chapter.content:
            self._render_content(doc, chapter.content)
            doc.add_paragraph()

        # 第二章：企业简介 + 资质清单
        if "企业" in chapter.title or "资格" in chapter.title:
            if enterprise:
                self._render_enterprise_info(doc, enterprise)

            if credentials:
                doc.add_paragraph()
                doc.add_heading("资质证书清单", level=2)
                self._render_credential_table(doc, credentials)

        # 第九章：业绩案例
        if "业绩" in chapter.title or "荣誉" in chapter.title:
            perf_creds = [c for c in credentials if c.cred_type in ("performance", "award")]
            if perf_creds:
                doc.add_heading("业绩及荣誉清单", level=2)
                self._render_credential_table(doc, perf_creds)

    def _render_enterprise_info(self, doc: Document, ent: Enterprise):
        """渲染企业基本信息表格"""
        doc.add_heading("企业基本信息", level=2)

        rows_data = [
            ("企业名称", ent.name),
            ("统一社会信用代码", ent.credit_code or "—"),
            ("法定代表人", ent.legal_representative or "—"),
            ("注册资本", f"{ent.registered_capital}万元" if ent.registered_capital else "—"),
            ("成立日期", ent.established_date or "—"),
            ("员工人数", f"{ent.employee_count}人" if ent.employee_count else "—"),
            ("联系地址", ent.address or "—"),
            ("联系人", ent.contact_person or "—"),
            ("联系电话", ent.contact_phone or "—"),
        ]

        # 食品专属
        if ent.food_license_no:
            rows_data.append(("食品经营许可证号", ent.food_license_no))

        certs = []
        if ent.haccp_certified:
            certs.append("HACCP")
        if ent.iso22000_certified:
            certs.append("ISO22000")
        if ent.sc_certified:
            certs.append("SC")
        if certs:
            rows_data.append(("体系认证", " / ".join(certs)))

        # 冷链资产
        if ent.cold_chain_vehicles:
            rows_data.append(("冷链车辆", f"{ent.cold_chain_vehicles}辆"))
        if ent.normal_vehicles:
            rows_data.append(("常温车辆", f"{ent.normal_vehicles}辆"))
        if ent.warehouse_area:
            rows_data.append(("仓储面积", f"{ent.warehouse_area}㎡"))
        if ent.cold_storage_area:
            rows_data.append(("冷库面积", f"{ent.cold_storage_area}㎡"))

        tbl = doc.add_table(rows=len(rows_data), cols=2, style="Table Grid")
        tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
        for i, (label, value) in enumerate(rows_data):
            cell_label = tbl.rows[i].cells[0]
            cell_value = tbl.rows[i].cells[1]
            cell_label.text = label
            cell_value.text = str(value)
            # 标签列加粗
            for run in cell_label.paragraphs[0].runs:
                run.font.bold = True
                run.font.size = Pt(11)
            for run in cell_value.paragraphs[0].runs:
                run.font.size = Pt(11)

        # 企业简介
        if ent.description:
            doc.add_paragraph()
            doc.add_heading("企业简介", level=2)
            self._render_content(doc, ent.description)

        # 竞争优势
        if ent.competitive_advantages:
            doc.add_paragraph()
            doc.add_heading("核心竞争优势", level=2)
            self._render_content(doc, ent.competitive_advantages)

    def _render_credential_table(self, doc: Document, credentials: list[Credential]):
        """渲染资质证书表格"""
        headers = ["序号", "证书类型", "证书名称", "证书编号", "有效期", "发证机关"]
        tbl = doc.add_table(rows=1, cols=len(headers), style="Table Grid")
        tbl.alignment = WD_TABLE_ALIGNMENT.CENTER

        # 表头
        for j, h in enumerate(headers):
            cell = tbl.rows[0].cells[j]
            cell.text = h
            for run in cell.paragraphs[0].runs:
                run.font.bold = True
                run.font.size = Pt(10)
            cell.paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

        # 数据行
        for i, cred in enumerate(credentials, 1):
            row = tbl.add_row()
            row.cells[0].text = str(i)
            row.cells[1].text = _CRED_TYPE_LABELS.get(cred.cred_type, cred.cred_type)
            row.cells[2].text = cred.cred_name
            row.cells[3].text = cred.cred_no or "—"

            if cred.is_permanent:
                row.cells[4].text = "长期有效"
            elif cred.expiry_date:
                row.cells[4].text = f"至 {cred.expiry_date}"
            else:
                row.cells[4].text = "—"

            row.cells[5].text = cred.issuing_authority or "—"

            for cell in row.cells:
                cell.paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                for run in cell.paragraphs[0].runs:
                    run.font.size = Pt(10)

    def _render_quotation_chapter(
        self, doc: Document, chapter: BidChapter, quotation: QuotationSheet
    ):
        """渲染报价文件章节（含报价表格）"""
        # 先输出章节内容（如有）
        if chapter.content:
            self._render_content(doc, chapter.content)
            doc.add_paragraph()

        # 报价汇总
        doc.add_heading("报价汇总", level=2)
        summary_data = []
        if quotation.budget_amount:
            summary_data.append(("招标预算金额", f"¥{quotation.budget_amount:,.2f}"))
        if quotation.total_amount:
            summary_data.append(("投标报价总额", f"¥{quotation.total_amount:,.2f}"))
        if quotation.discount_rate:
            summary_data.append(("下浮率", f"{quotation.discount_rate * 100:.1f}%"))
        if quotation.pricing_method:
            methods = {"fixed_price": "固定单价", "discount_rate": "下浮率", "comprehensive": "综合报价"}
            summary_data.append(("报价方式", methods.get(quotation.pricing_method, quotation.pricing_method)))

        if summary_data:
            tbl = doc.add_table(rows=len(summary_data), cols=2, style="Table Grid")
            tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
            for i, (label, value) in enumerate(summary_data):
                tbl.rows[i].cells[0].text = label
                tbl.rows[i].cells[1].text = value
                for run in tbl.rows[i].cells[0].paragraphs[0].runs:
                    run.font.bold = True
                    run.font.size = Pt(11)
                for run in tbl.rows[i].cells[1].paragraphs[0].runs:
                    run.font.size = Pt(11)

        # 报价明细表
        if quotation.items:
            doc.add_paragraph()
            doc.add_heading("报价明细表", level=2)

            headers = ["序号", "品类", "品名", "规格", "单位", "单价(元)", "数量", "小计(元)"]
            tbl = doc.add_table(rows=1, cols=len(headers), style="Table Grid")
            tbl.alignment = WD_TABLE_ALIGNMENT.CENTER

            for j, h in enumerate(headers):
                cell = tbl.rows[0].cells[j]
                cell.text = h
                for run in cell.paragraphs[0].runs:
                    run.font.bold = True
                    run.font.size = Pt(9)
                cell.paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

            for i, item in enumerate(quotation.items, 1):
                row = tbl.add_row()
                row.cells[0].text = str(i)
                row.cells[1].text = _FOOD_CATEGORY_LABELS.get(item.category, item.category)
                row.cells[2].text = item.item_name
                row.cells[3].text = item.spec or "—"
                row.cells[4].text = item.unit or "—"
                row.cells[5].text = f"{item.unit_price:.2f}" if item.unit_price else "—"
                row.cells[6].text = f"{item.quantity:.1f}" if item.quantity else "—"
                row.cells[7].text = f"{item.amount:.2f}" if item.amount else "—"

                for cell in row.cells:
                    cell.paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                    for run in cell.paragraphs[0].runs:
                        run.font.size = Pt(9)

        if quotation.remarks:
            doc.add_paragraph()
            p = doc.add_paragraph()
            run = p.add_run(f"备注：{quotation.remarks}")
            run.font.size = Pt(11)
            run.font.italic = True

    def _render_content(self, doc: Document, content: str):
        """智能排版引擎：将章节正文渲染为格式化 Word 段落"""
        for line in content.split("\n"):
            stripped = line.strip()
            if not stripped:
                continue

            # Markdown 标题 → Word heading
            heading_match = _HEADING_RE.match(stripped)
            if heading_match:
                level = len(stripped) - len(stripped.lstrip("#"))
                text = stripped.lstrip("# ").strip()
                doc.add_heading(text, level=min(level + 1, 4))
                continue

            # 警告行
            if stripped.startswith("⚠") or stripped.startswith("**⚠"):
                p = doc.add_paragraph()
                run = p.add_run(stripped.replace("**", ""))
                run.font.color.rgb = RGBColor(0xCC, 0x00, 0x00)
                run.font.bold = True
                continue

            # 节标题
            if _SECTION_TITLE_RE.match(stripped):
                doc.add_heading(stripped, level=3)
                continue

            # 大编号（一、二、三、）
            if _NUMBERED_ITEM_RE.match(stripped):
                p = doc.add_paragraph()
                run = p.add_run(stripped)
                run.font.size = Pt(12)
                run.font.bold = True
                continue

            # 条标题
            if _CLAUSE_TITLE_RE.match(stripped):
                p = doc.add_paragraph()
                run = p.add_run(stripped)
                run.font.size = Pt(12)
                run.font.bold = True
                continue

            # 数字子项 1. 2. （一）
            if _SUB_ITEM_RE.match(stripped):
                p = doc.add_paragraph(stripped)
                pf = p.paragraph_format
                pf.left_indent = Pt(12)
                pf.space_before = Pt(2)
                for run in p.runs:
                    run.font.size = Pt(12)
                continue

            # 款项 （1） ①
            if _SUB_CLAUSE_RE.match(stripped):
                p = doc.add_paragraph(stripped)
                pf = p.paragraph_format
                pf.left_indent = Pt(24)
                pf.space_before = Pt(3)
                for run in p.runs:
                    run.font.size = Pt(12)
                continue

            # 【】加粗标记
            if stripped.startswith("【"):
                p = doc.add_paragraph()
                run = p.add_run(stripped)
                run.font.bold = True
                run.font.size = Pt(12)
                continue

            # Markdown 表格行 → 跳过分隔线，渲染数据行为段落
            if stripped.startswith("|") and stripped.endswith("|"):
                if re.match(r'^\|[\s\-:]+\|', stripped):
                    continue  # 分隔线跳过
                p = doc.add_paragraph(stripped)
                for run in p.runs:
                    run.font.size = Pt(10)
                continue

            # 普通正文
            p = doc.add_paragraph(stripped)
            pf = p.paragraph_format
            pf.first_line_indent = Pt(24)
            pf.space_before = Pt(2)
            pf.space_after = Pt(2)
            pf.line_spacing = Pt(22)
