import asyncio
"""

流程：
  1. 加载 ProjectParams + Project 基础信息
  2. 调用 RuleService.match_rules() → 命中规则+章节列表
  3. 调用 SupportCalcEngine + VentCalcEngine → 计算结果
  4. 按章节顺序组装内容
  5. python-docx 生成 .docx 文件
"""
import os
from datetime import datetime
from typing import Optional

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

from sqlalchemy import select
from sqlalchemy.ext.asyncio import AsyncSession

from app.models.project import Project, ProjectParams
from app.schemas.calc import SupportCalcInput
from app.schemas.vent import VentCalcInput
from app.schemas.doc import ChapterContent, DocGenerateResult
from app.services.calc_engine import SupportCalcEngine
from app.services.vent_engine import VentCalcEngine
from app.services.rule_service import RuleService


# 参数字段中文映射
PARAM_LABELS: dict[str, str] = {
    "rock_class": "围岩级别", "coal_thickness": "煤层厚度(m)",
    "coal_dip_angle": "煤层倾角(°)", "gas_level": "瓦斯等级",
    "hydro_type": "水文地质类型", "geo_structure": "地质构造",
    "spontaneous_combustion": "自燃倾向性", "roadway_type": "巷道类型",
    "excavation_type": "掘进类型", "section_form": "断面形式",
    "section_width": "断面宽度(m)", "section_height": "断面高度(m)",
    "excavation_length": "掘进长度(m)", "service_years": "服务年限(年)",
    "dig_method": "掘进方式", "dig_equipment": "掘进设备",
    "transport_method": "运输方式",
}


class DocGenerator:
    """文档生成引擎"""

    def __init__(self, session: AsyncSession):
        self.session = session

    async def generate(
        self, project_id: int, tenant_id: int, include_calc: bool = True
    ) -> DocGenerateResult:
        """
        端到端文档生成

        Returns:
            DocGenerateResult 包含文件路径和章节列表
        """
        # 1. 加载项目信息
        project = await self._load_project(project_id, tenant_id)
        if not project:
            raise ValueError("项目不存在或无权访问")

        params = await self._load_params(project_id)
        params_dict = self._params_to_dict(params) if params else {}

        # 2. 规则匹配
        rule_svc = RuleService(self.session)
        match_result = None
        try:
            match_result = await rule_svc.match_rules(project_id, tenant_id)
        except ValueError:
            pass  # 参数未填写时跳过匹配

        # 3. 计算引擎
        calc_result = None
        vent_result = None
        if include_calc and params:
            calc_result = self._run_support_calc(params_dict)
            vent_result = self._run_vent_calc(params_dict)

        # 4. 组装章节
        chapters = self._assemble_chapters(
            project, params_dict, match_result, calc_result, vent_result
        )

        # 4.5 智能深度润色（AI赋能）
        chapters = await self._ai_polish_content(chapters, params_dict)

        # 5. 生成 Word
        file_path = self._render_docx(project, chapters, calc_result, vent_result)

        total_warnings = 0
        if calc_result:
            total_warnings += len(calc_result.warnings)
        if vent_result:
            total_warnings += len(vent_result.warnings)

        return DocGenerateResult(
            project_id=project_id,
            project_name=project.face_name,
            file_path=file_path,
            total_chapters=len(chapters),
            total_warnings=total_warnings,
            chapters=chapters,
        )

    # ========== 加载数据 ==========

    async def _load_project(self, pid: int, tid: int) -> Optional[Project]:
        result = await self.session.execute(
            select(Project).where(Project.id == pid, Project.tenant_id == tid)
        )
        return result.scalar_one_or_none()

    async def _load_params(self, pid: int) -> Optional[ProjectParams]:
        result = await self.session.execute(
            select(ProjectParams).where(ProjectParams.project_id == pid)
        )
        return result.scalar_one_or_none()

    @staticmethod
    def _params_to_dict(params: ProjectParams) -> dict:
        return {
            c.key: getattr(params, c.key)
            for c in params.__table__.columns
            if c.key not in ("id", "project_id")
        }

    # ========== 计算引擎调用 ==========

    @staticmethod
    def _run_support_calc(p: dict):
        try:
            inp = SupportCalcInput(
                rock_class=p.get("rock_class", "III"),
                section_form=p.get("section_form", "矩形"),
                section_width=float(p.get("section_width", 4.5)),
                section_height=float(p.get("section_height", 3.2)),
            )
            return SupportCalcEngine.calculate(inp)
        except Exception:
            return None

    @staticmethod
    def _run_vent_calc(p: dict):
        try:
            inp = VentCalcInput(
                gas_emission=float(p.get("gas_emission", 2.0) or 2.0),
                gas_level=p.get("gas_level", "低瓦斯"),
                section_area=float(p.get("section_width", 4.5)) * float(p.get("section_height", 3.2)),
                excavation_length=float(p.get("excavation_length", 300)),
            )
            return VentCalcEngine.calculate(inp)
        except Exception:
            return None

    # ========== 章节组装 ==========

    def _assemble_chapters(self, project, params_dict, match_result, calc_result, vent_result):
        """
        严格按华阳集团《采掘运技术管理规定》附件2
        《掘进工作面作业规程编制内容及提纲》组装规程

        官方9章结构:
          第一章 概述
          第二章 地面相对位置及水文地质概况 (4节)
          第三章 巷道布置及支护说明 (5节)
          第四章 施工工艺 (3节)
          第五章 生产系统 (7节)
          第六章 劳动组织及主要技术经济指标 (2节)
          第七章 煤质管理
          第八章 安全技术措施 (9节)
          第九章 安全风险管控及应急避险
          附录
        """
        chapters: list[ChapterContent] = []

        mine_name = getattr(project, 'mine_name', '—')
        roadway_type = params_dict.get('roadway_type', '—')
        gas_level = params_dict.get('gas_level', '—')
        excavation_type = params_dict.get('excavation_type', '—')

        # ================================================================
        #  第一章  概述
        # ================================================================
        ch1_lines = [
            f"一、巷道名称：{project.face_name}",
            f"  矿井名称：{mine_name}",
            f"  巷道用途/性质：{roadway_type}",
            f"  设计长度：{params_dict.get('excavation_length', '—')} m",
            f"  坡度：按设计确定",
            "",
            "二、特殊技术要求及需要重点说明的问题",
            f"  掘进类型：{excavation_type}",
            f"  瓦斯等级：{gas_level}",
            f"  自燃倾向性：{params_dict.get('spontaneous_combustion', '—')}",
            f"  掘进方式：{params_dict.get('dig_method', '—')}",
            "",
            "三、巷道布置平面图",
            "  （详见附图）",
        ]
        chapters.append(ChapterContent(
            chapter_no="第一章", title="概述",
            content="\n".join(ch1_lines), source="template",
        ))

        # ================================================================
        #  第二章  地面相对位置及水文地质概况（4节）
        # ================================================================
        ch2_lines = [
            "第一节  地面相对位置及邻近采区开采情况",
            "",
            f"一、巷道相应的地面位置、标高，区域内的水体和建筑物对工程的影响",
            f"  （根据矿井实际填写）",
            "",
            "二、巷道与相邻煤（岩）层、邻近巷道的层间关系",
            "  附近已有的采掘情况对工程的影响。",
            "",
            "三、老空区的水、火、瓦斯等对工程的影响分析",
            "",
            "第二节  煤（岩）层赋存特征",
            "",
            "一、煤（岩）层产状、厚度、结构",
            f"  煤层厚度：{params_dict.get('coal_thickness', '—')} m",
            f"  煤层倾角：{params_dict.get('coal_dip_angle', '—')}°",
            f"  坚固性系数(f)：根据实测确定",
            "",
            "二、预测巷道瓦斯涌出量、煤层自然发火倾向等",
            f"  瓦斯等级：{gas_level}",
            f"  自燃倾向性：{params_dict.get('spontaneous_combustion', '—')}",
            "",
            "三、其他煤（岩）层技术特征分析",
            "",
            "四、地层综合柱状图（详见附图）",
            "",
            "五、巷道围岩类别",
            f"  围岩级别：{params_dict.get('rock_class', '—')}",
            "",
            "第三节  地质构造",
            "",
            f"一、巷道煤（岩）层产状，断层、褶曲等地质构造要素",
            f"  地质构造特征：{params_dict.get('geo_structure', '—')}",
            "",
            "二、地质平面图、剖面图（详见附图）",
            "",
            "第四节  水文地质",
            "",
            "一、主要充水因素分析",
            f"  水文地质类型：{params_dict.get('hydro_type', '—')}",
            "",
            "二、带压掘进工作面突水系数计算及危险性评价",
            "",
            "三、积水区域附近掘进巷道，标出\"三线\"（积水线、探水线和警戒线）",
            "",
            "四、预测工作面正常、最大涌水量",
        ]
        chapters.append(ChapterContent(
            chapter_no="第二章", title="地面相对位置及水文地质概况",
            content="\n".join(ch2_lines), source="template",
        ))

        # ================================================================
        #  第三章  巷道布置及支护说明（5节）
        # ================================================================
        sec_w = float(params_dict.get("section_width", 0) or 0)
        sec_h = float(params_dict.get("section_height", 0) or 0)
        sec_area = round(sec_w * sec_h, 2) if sec_w > 0 and sec_h > 0 else 0

        ch3_lines = [
            "第一节  巷道布置",
            "",
            f"一、巷道布置：层位、水平标高、开口的位置、方位角",
            f"  巷道类型：{roadway_type}",
            f"  掘进长度：{params_dict.get('excavation_length', '—')} m",
            f"  服务年限：{params_dict.get('service_years', '—')} 年",
            "",
            "二、特殊地点的施工（车场、硐室、交岔点等）",
            "",
            "三、开口大样图（详见附图）",
            "",
            "第二节  矿压观测",
            "",
            "一、采用锚网支护掘进巷道必须安设顶板离层仪和锚杆锚索测力计。",
            "",
            "二、矿压观测分综合监测和一般监测两种：",
            "  综合监测：巷道表面位移、顶板离层、锚杆锚索受力状况监测。",
            "  一般监测：巷道表面位移、顶板离层监测。",
            "",
            "三、回采顺槽巷道设置综合测站和一般测站，明确布站间距、数量及观测分析标准。",
            "",
            "四、特殊地段（开口处、交岔点、构造影响区等）须增设监测仪器。",
            "",
            "第三节  顶板岩性探测",
            "",
            "一、掘进巷道必须进行顶板岩性探测与分析，验证与优化支护设计。",
            "",
            "二、明确岩性探测方法和具体技术要求。",
            "",
            "第四节  支护设计",
            "",
            f"一、巷道断面设计",
            f"  断面形式：{params_dict.get('section_form', '—')}",
            f"  断面宽度：{sec_w} m" if sec_w > 0 else "  断面宽度：—",
            f"  断面高度：{sec_h} m" if sec_h > 0 else "  断面高度：—",
            f"  断面净面积：{sec_area} m²" if sec_area > 0 else "  断面净面积：—",
            "",
        ]

        # 嵌入支护计算结果
        if calc_result:
            ch3_lines.extend([
                "二、支护参数（计算引擎输出）",
                f"  单根锚杆锚固力：{calc_result.bolt_force} kN",
                f"  最大允许锚杆间距：{calc_result.max_bolt_spacing} mm",
                f"  最大允许排距：{calc_result.max_bolt_row_spacing} mm",
                f"  推荐每排锚杆数：{calc_result.recommended_bolt_count_per_row} 根",
                f"  最少锚索数量：{calc_result.min_cable_count} 根",
                f"  支护密度：{calc_result.support_density} 根/m²",
                f"  安全系数：{calc_result.safety_factor}",
                "",
            ])
            if calc_result.warnings:
                ch3_lines.append("【支护合规预警】")
                for w in calc_result.warnings:
                    ch3_lines.append(f"  ⚠ {w.message}")
                ch3_lines.append("")
        else:
            ch3_lines.extend([
                "二、支护参数",
                "  （待填写支护设计参数）",
                "",
            ])

        ch3_lines.extend([
            "三、支护参数校核",
            "  （一）顶锚杆校核：L ≥ L1 + L2 + L3",
            "  （二）校核顶锚杆间排距",
            "  （三）加强锚索长度校核",
            "  （四）加强锚索数目校核",
            "",
            "四、巷道断面图、平面图、交岔点支护示意图（详见附图）",
            "",
            "五、支护设计采用动态信息设计方法",
            "  工程类比法初始设计→持续矿压观测→修改优化→正式设计。",
            "  巷道条件发生变化时，须立即组织现场查看并及时修改支护设计。",
            "",
            "第五节  支护工艺",
            "",
            "一、临时支护工艺、工序及要求",
            "  煤巷综掘、大断面岩巷综掘必须使用机载式临时支护装置。",
            "  临时支护与永久支护距掘进工作面的距离须在规程中明确。",
            "",
            "二、永久支护工艺、工序及要求",
            "  （一）锚杆及联合支护：材质、规格、间排距、安装、锚固力要求。",
            "  （二）支架支护：构件齐全，背紧背牢、充满填实。",
            "",
            "三、施工质量标准表（详见附表）",
        ])

        chapters.append(ChapterContent(
            chapter_no="第三章", title="巷道布置及支护说明",
            content="\n".join(ch3_lines), source="calc_engine" if calc_result else "template",
            has_warning=bool(calc_result and not calc_result.is_compliant),
        ))

        # ================================================================
        #  第四章  施工工艺（3节）
        # ================================================================
        ch4_lines = [
            "第一节  施工方法",
            "",
            "一、确定巷道施工方法",
            f"  掘进方式：{params_dict.get('dig_method', '—')}",
            f"  掘进类型：{excavation_type}",
            "",
            "二、巷道开口施工方法",
            "  从支设临时支护开始，到永久支护止的施工顺序。",
            "",
            "三、特殊条件下的施工方法",
            "  （一）石门揭开煤层时：打超前钻排放瓦斯、远距离放炮。",
            "  （二）硐室的施工方法：根据围岩类别选用全断面或分层施工法。",
            "  （三）交岔点的施工方法：根据围岩类别选用相应施工法。",
            "  （四）倾斜巷道：支架迎山角、防滑防跑车装置。",
            "",
            "第二节  掘进方式",
            "",
            f"一、掘进方式：{params_dict.get('dig_method', '—')}",
            f"  掘进设备：{params_dict.get('dig_equipment', '—')}",
            "",
            "二、机掘作业方式、截割顺序、截割循环进度",
            "",
            "三、炮掘施工工序安排、工艺流程",
            "  严格执行\"一炮三检\"和\"三人连锁放炮\"制度。",
            "  炮孔内发现异状、温度骤高骤低、有显著瓦斯涌出、煤岩松软时，",
            "  必须停止装药，并采取安全措施。",
            "",
            "第三节  装载运输",
            "",
            f"一、运输方式：{params_dict.get('transport_method', '—')}",
            "",
            "二、装载设备及操作要求",
            "",
            "三、管线及轨道敷设",
            "  风、水管路要同侧敷设，不得与瓦斯抽采管路同侧布置。",
        ]

        chapters.append(ChapterContent(
            chapter_no="第四章", title="施工工艺",
            content="\n".join(ch4_lines), source="template",
        ))

        # ================================================================
        #  第五章  生产系统（7节）
        # ================================================================
        ch5_lines = [
            "第一节  一通三防",
            "",
            "一、通风系统",
        ]
        if vent_result:
            ch5_lines.extend([
                f"  瓦斯涌出法需风量(Q_gas)：{vent_result.q_gas} m³/min",
                f"  人数法需风量(Q_people)：{vent_result.q_people} m³/min",
                f"  炸药法需风量(Q_explosive)：{vent_result.q_explosive} m³/min",
                f"  最终配风量(Q_required)：{vent_result.q_required} m³/min",
                f"  推荐局扇型号：{vent_result.recommended_fan}（{vent_result.fan_power} kW）",
            ])
            if vent_result.warnings:
                ch5_lines.append("  【通风合规预警】")
                for w in vent_result.warnings:
                    ch5_lines.append(f"    ⚠ {w.message}")
        else:
            ch5_lines.append("  （待计算通风参数）")

        ch5_lines.extend([
            "",
            "二、综合防尘",
            "  采用湿式打眼、喷雾降尘、通风除尘等综合防尘措施。",
            "  掘进工作面应安设净化水幕、转载点喷雾装置。",
            "",
            "三、防灭火",
            "  巷道布置及通风设施设置须符合防灭火要求。",
            "  厚煤层工作面进、回风巷开口须按规定设置防灭火设施。",
            "",
            "第二节  压风",
            "",
            "  压风自救装置安设间距不超过200米。",
            "  压风管路规格、连接方式及维护要求。",
            "",
            "第三节  动力（供电）",
            "",
            "  掘进工作面供电须采用三专线路（专用变压器、开关、电缆）。",
            "  电气设备选型及防爆要求。",
            "",
            "第四节  排水",
            "",
            "  掘进工作面须配备排水设备，排水能力须满足最大涌水量要求。",
            f"  水文地质类型：{params_dict.get('hydro_type', '—')}",
            "",
            "第五节  运输",
            "",
            f"  运输方式：{params_dict.get('transport_method', '—')}",
            "  运输设备选型、安全措施。",
            "",
            "第六节  通讯照明",
            "",
            "  掘进工作面应设置有线调度电话和应急通讯设备。",
            "  照明灯具选型及安设要求。",
            "",
            "第七节  供水施救",
            "",
            "  供水施救管路安装要求。",
            "  每隔一定距离设置三通阀门。",
        ])

        chapters.append(ChapterContent(
            chapter_no="第五章", title="生产系统",
            content="\n".join(ch5_lines), source="calc_engine" if vent_result else "template",
            has_warning=bool(vent_result and not vent_result.is_compliant),
        ))

        # ================================================================
        #  第六章  劳动组织及主要技术经济指标（2节）
        # ================================================================
        ch6_lines = [
            "第一节  劳动组织",
            "",
            "一、劳动组织",
            "  实行\"三八\"作业制，两班生产、一班检修。",
            "",
            "二、正规循环作业",
            "  按照\"正规循环作业\"要求组织生产，正规循环率不低于80%。",
            "",
            "三、循环作业图表（详见附表）",
            "",
            "第二节  主要技术经济指标",
            "",
            f"  巷道掘进长度：{params_dict.get('excavation_length', '—')} m",
            f"  断面形式：{params_dict.get('section_form', '—')}",
            f"  断面净面积：{sec_area} m²" if sec_area > 0 else "  断面净面积：—",
            "  循环进度、月进度、工效等（详见技术经济指标表）",
        ]

        chapters.append(ChapterContent(
            chapter_no="第六章", title="劳动组织及主要技术经济指标",
            content="\n".join(ch6_lines), source="template",
        ))

        # ================================================================
        #  第七章  煤质管理（官方新增章节）
        # ================================================================
        ch7_lines = [
            "一、煤质指标",
            "  简要说明煤质指标：灰分、水分、发热量、硫分等。",
            "",
            "二、提高煤质及采出率的措施",
            "  （一）严格控制混矸率，掘进中分层排矸。",
            "  （二）煤岩分装分运，防止煤矸混装。",
            "  （三）减少煤尘飞扬损失。",
            "  （四）合理确定巷道断面，提高煤炭采出率。",
        ]

        chapters.append(ChapterContent(
            chapter_no="第七章", title="煤质管理",
            content="\n".join(ch7_lines), source="template",
        ))

        # ================================================================
        #  第八章  安全技术措施（9节）
        # ================================================================
        ch8_lines = [
            "第一节  一般规定",
            "",
            "一、工作面安全管理",
            "  有针对性地叙述与本工作面相关的安全制度及需特别强调的安全措施。",
            "",
            "二、交接班管理",
            "  叙述交接班安全检查内容和有关规定。",
            "",
            "第二节  顶板管理",
            "",
            "一、掘进工作面应当严格执行敲帮问顶制度，开工前必须全面检查。",
            "二、空顶距离不得超过规定值，掘进面附近必须设临时支护。",
            "三、构造影响段的顶板管理措施。",
            "四、离层仪和锚杆锚索测力计的监测要求。",
            "",
            "第三节  一通三防",
            "",
            f"一、本工作面瓦斯等级：{gas_level}",
            "二、初揭煤层、过构造等特殊时期的瓦斯防治措施。",
            "三、工作面通风路线发生风流不畅情况下的应急处理措施。",
            "四、各类综合防尘设施的使用管理要求。",
            f"五、{'煤与瓦斯突出工作面防突施工要求及安全注意事项。' if gas_level == '突出' else '瓦斯监测要求。'}",
            "六、工作面火灾的预防及处置。",
            "",
            "第四节  爆破",
            "",
            "一、一般规定：瓦检员、放炮员等职责分工。",
            "二、爆破施工要求：炮眼布置图、说明表。",
            "三、爆破作业规定：",
            "  严格执行\"一炮三检\"和\"三人连锁放炮\"制度。",
            "  瓦斯浓度达到0.5%时严禁放炮。",
            "",
            "第五节  机电",
            "",
            "一、一般规定：操作专职制、设备检修维护。",
            "二、机械设备管理：操作使用、维修更换安全措施。",
            "三、电气设备管理：",
            "  井下电气设备必须取得煤矿矿用产品安全标志。",
            "  各部位绝缘电阻值应不低于0.5MΩ。",
            "",
            "第六节  运输",
            "",
            "一、一般规定：运输设备运行前必须发出警报信号。",
            "二、主运输设备管理。",
            "三、辅助运输管理。",
            "",
            "第七节  监控与通讯",
            "",
            "一、掘进工作面须安设甲烷传感器、一氧化碳传感器、风速传感器。",
            f"  {'高瓦斯矿井使用量程上限不低于10%的甲烷传感器。' if gas_level in ('高瓦斯', '突出') else ''}",
            f"  {'突出煤层使用量程上限不低于40%的甲烷传感器。' if gas_level == '突出' else ''}",
            "二、井下作业人员佩戴、使用人员位置识别卡。",
            "三、通讯设备联络、使用、维护要求。",
            "四、井下视频监控的使用、维护要求。",
            "",
            "第八节  防治水",
            "",
            f"一、水文地质类型：{params_dict.get('hydro_type', '—')}",
            "二、坚持\"有疑必探、先探后掘\"的原则。",
            "三、防治水管控措施。",
            "四、透水征兆等异常情况时的应急措施。",
            "",
            "第九节  其他",
            "",
            "一、安全生产标准化重点内容及工程质量、文明生产相关标准要求。",
            "二、其他安全技术措施。",
        ]

        chapters.append(ChapterContent(
            chapter_no="第八章", title="安全技术措施",
            content="\n".join(ch8_lines), source="template",
        ))

        # ================================================================
        #  第九章  安全风险管控及应急避险
        # ================================================================
        ch9_lines = [
            "一、安全风险管控",
            "  按照安全风险专项辨识评估资料，重点描述工作面作业环境、",
            "  工程技术、设备设施、现场操作等方面存在的安全风险及相应管控措施。",
            f"  瓦斯等级：{gas_level}",
            f"  自燃倾向性：{params_dict.get('spontaneous_combustion', '—')}",
            f"  水文地质类型：{params_dict.get('hydro_type', '—')}",
            "",
            "二、紧急避险设施",
            "  掘进工作面所有作业人员必须随身携带自救器。",
            "  压风自救装置安设间距不超过200米。",
            "  明确压风及供水施救、避难硐室等设施的设置、规格及要求。",
            "",
            "三、灾害应急处置措施",
            "  （一）制定发生顶板、瓦斯、煤尘爆炸、火灾、水灾等事故的应急程序。",
            "  明确授予带班人员、班组长、瓦检工、调度人员遇险处置权和紧急避险权。",
            "  （二）确定发生灾害时的自救方式、组织抢救方法和安全撤离路线。",
            "  （三）明确避灾原则及安全注意事项。",
            "  绘制工作面避灾路线示意图（分水、火、瓦斯、煤尘路线标识）。",
            "  （详见附图：避灾路线图）",
        ]

        chapters.append(ChapterContent(
            chapter_no="第九章", title="安全风险管控及应急避险",
            content="\n".join(ch9_lines), source="template",
        ))

        # ================================================================
        #  附录：编制依据与规则命中
        # ================================================================
        if match_result and match_result.matched_rules:
            rule_lines = []
            for mr in match_result.matched_rules:
                rule_lines.append(f"• {mr.rule_name}（{mr.category}，优先级 {mr.priority}）")
                for a in mr.actions:
                    rule_lines.append(f"  → 关联章节：{a.target_chapter}")

            chapters.append(ChapterContent(
                chapter_no="附录", title="编制依据与规则命中",
                content="\n".join(rule_lines), source="rule_match",
            ))

        return chapters


    async def _ai_polish_content(self, chapters: list[ChapterContent], params: dict) -> list[ChapterContent]:
        """
        AI 深度润色长尾章节 — RAG 增强版

        流程:
          1. 对每个须润色的章节，用章节标题检索标准库 + 知识库
          2. 将检索到的规程片段注入 LLM System Prompt
          3. LLM 基于客户真实规程生成更贴合实际的内容
        """
        from app.core.config import settings
        from openai import AsyncOpenAI
        from app.services.embedding_service import EmbeddingService

        api_key = settings.OPENAI_API_KEY or settings.GEMINI_API_KEY
        base_url = settings.OPENAI_BASE_URL or None
        model = settings.AI_MODEL

        if not api_key:
            return chapters  # 降级：无大模型配置时原样返回

        client_kwargs = {"api_key": api_key}
        if base_url:
            client_kwargs["base_url"] = base_url
        client = AsyncOpenAI(**client_kwargs)

        # RAG 检索服务
        emb_svc = EmbeddingService(self.session)

        # 单章超时（秒）— 超时自动降级为模板内容
        POLISH_TIMEOUT = 30

        async def _polish_one(ch: ChapterContent) -> None:
            """单章 AI 润色 — 超时或异常时静默降级"""
            if not any(kw in ch.title for kw in [
                "安全技术措施", "支护", "应急", "施工工艺",
                "生产系统", "水文地质", "煤质管理", "安全风险",
            ]):
                return  # 不需要润色的章节直接跳过

            # ===== RAG 检索：标准库 + 知识库 =====
            rag_context_parts = []

            # L1a: 标准库条款
            try:
                std_results = await emb_svc.search_similar(
                    query=ch.title, tenant_id=1, top_k=3, threshold=0.4
                )
                if std_results:
                    rag_context_parts.append("【标准库参考条款】")
                    for r in std_results:
                        rag_context_parts.append(
                            f"- [{r['doc_title']}] {r['clause_no']}: {r['content'][:300]}"
                        )
            except Exception:
                pass  # RAG 检索失败不影响生成

            # L1b: 知识库（客户规程片段）
            try:
                snippet_results = await emb_svc.search_snippets(
                    query=ch.title, tenant_id=1, top_k=5, threshold=0.4
                )
                if snippet_results:
                    rag_context_parts.append("\n【客户规程参考内容】")
                    for r in snippet_results:
                        rag_context_parts.append(
                            f"- [{r['chapter_name']}]: {r['content'][:300]}"
                        )
            except Exception:
                pass

            rag_context = "\n".join(rag_context_parts)
            rag_note = ""
            if rag_context:
                rag_note = (
                    f"\n\n以下是从客户已有规程和国家标准中检索到的相关内容，"
                    f"请务必参考并融入你的输出中，确保与客户实际情况一致：\n{rag_context}\n"
                )

            prompt = (
                f"请作为顶尖煤矿安全专家，根据以下参数对作业规程的【{ch.title}】章节进行扩充、润色，"
                f"使其更符合现场实际，具备可操作性，避免生硬的模板拼接感。\n"
                f"地质条件与参数: {params}\n"
                f"原始内容框架:\n{ch.content}"
                f"{rag_note}\n\n"
                "要求：\n"
                "1. 直接输出润色后的篇章正式内容，不要包含任何前言后语和分析推理过程。\n"
                "2. 分条列出，层级清晰，重点突出。\n"
                "3. 如引用了参考资料中的具体数值标准，请在文中自然融入。"
            )
            try:
                resp = await asyncio.wait_for(
                    client.chat.completions.create(
                        model=model,
                        messages=[{"role": "user", "content": prompt}]
                    ),
                    timeout=POLISH_TIMEOUT,
                )
                polished = resp.choices[0].message.content
                if polished:
                    ch.content = polished
                    ch.source = "ai_polished"
            except asyncio.TimeoutError:
                print(f"⏱️ AI 润色超时({POLISH_TIMEOUT}s), 降级使用模板: {ch.title}")
            except Exception as e:
                print(f"⚠️ AI 润色失败: {ch.title}: {e}")

        # 并发润色所有章节（从串行 ~150s 降至并发 ~30s）
        await asyncio.gather(*[_polish_one(ch) for ch in chapters])

        return chapters

    # ========== Word 渲染 ==========

    def _render_docx(self, project, chapters, calc_result, vent_result) -> str:
        """用 python-docx 生成 .docx 文件"""
        doc = Document()

        # 文档样式
        style = doc.styles["Normal"]
        style.font.name = "宋体"
        style.font.size = Pt(12)

        # --- 封面 ---
        for _ in range(4):
            doc.add_paragraph()

        title = doc.add_paragraph()
        title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        run = title.add_run(f"{project.face_name}")
        run.font.size = Pt(22)
        run.font.bold = True

        subtitle = doc.add_paragraph()
        subtitle.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        run2 = subtitle.add_run("掘进工作面作业规程")
        run2.font.size = Pt(18)

        doc.add_paragraph()

        meta_items = [
            f"矿井名称：{getattr(project, 'mine_name', '—')}",
            f"编制日期：{datetime.now().strftime('%Y年%m月%d日')}",
            f"编制单位：生产技术科",
        ]
        for item in meta_items:
            p = doc.add_paragraph()
            p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            p.add_run(item).font.size = Pt(14)

        doc.add_page_break()

        # --- 正文章节 ---
        for ch in chapters:
            # 章节标题
            heading = doc.add_heading(f"{ch.chapter_no}  {ch.title}", level=1)

            # 预警标记
            if ch.has_warning:
                warn_p = doc.add_paragraph()
                warn_run = warn_p.add_run("⚠ 本章节存在合规预警，请重点审查")
                warn_run.font.color.rgb = RGBColor(0xCC, 0x00, 0x00)
                warn_run.font.bold = True

            # 章节内容
            for line in ch.content.split("\n"):
                if line.startswith("  ⚠"):
                    p = doc.add_paragraph()
                    run = p.add_run(line)
                    run.font.color.rgb = RGBColor(0xCC, 0x00, 0x00)
                elif line.startswith("【"):
                    p = doc.add_paragraph()
                    run = p.add_run(line)
                    run.font.bold = True
                else:
                    doc.add_paragraph(line)

        # --- 保存 ---
        output_dir = os.path.join(
            os.path.dirname(os.path.dirname(os.path.dirname(os.path.abspath(__file__)))),
            "storage", "outputs"
        )
        os.makedirs(output_dir, exist_ok=True)

        safe_name = project.face_name.replace("/", "_").replace(" ", "_")
        filename = f"{safe_name}_作业规程_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
        file_path = os.path.join(output_dir, filename)

        doc.save(file_path)
        return file_path
